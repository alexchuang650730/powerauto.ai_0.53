#!/usr/bin/env python3
"""
Word文档生成器模块
将OCR识别和Gemini修正后的结果转换为Word文档
"""

import os
import sys
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import json
import time
from datetime import datetime

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class WordGenerator:
    """Word文档生成器，将OCR结果转换为Word文档"""
    
    def __init__(self, use_mock: bool = True):
        """
        初始化Word文档生成器
        
        Args:
            use_mock: 是否使用模拟模式（用于测试，不需要实际安装python-docx）
        """
        self.use_mock = use_mock
        self.docx_available = False
        
        # 如果不是模拟模式，则尝试导入python-docx
        if not self.use_mock:
            try:
                import docx
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                self.docx = docx
                self.Pt = Pt
                self.Inches = Inches
                self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
                self.docx_available = True
                logger.info("python-docx库加载成功")
            except ImportError:
                logger.warning("未找到python-docx库，将使用模拟模式")
                self.use_mock = True
        
        if self.use_mock:
            logger.info("使用模拟模式，不会生成实际的Word文档")
    
    def generate_document(self, ocr_result: Dict[str, Any], output_path: str) -> str:
        """
        根据OCR结果生成Word文档
        
        Args:
            ocr_result: OCR处理结果（可能已经过Gemini修正）
            output_path: 输出文档路径
            
        Returns:
            生成的文档路径
        """
        start_time = time.time()
        logger.info(f"开始生成Word文档: {output_path}")
        
        # 检查OCR结果是否有效
        if "error" in ocr_result:
            logger.error(f"OCR结果包含错误: {ocr_result['error']}")
            return ""
        
        # 如果是模拟模式，创建一个简单的文本文件
        if self.use_mock:
            return self._generate_mock_document(ocr_result, output_path)
        
        # 使用python-docx生成实际的Word文档
        if self.docx_available:
            return self._generate_real_document(ocr_result, output_path)
        else:
            logger.error("python-docx库不可用，无法生成Word文档")
            return ""
    
    def _generate_mock_document(self, ocr_result: Dict[str, Any], output_path: str) -> str:
        """
        生成模拟的Word文档（实际上是文本文件）
        
        Args:
            ocr_result: OCR处理结果
            output_path: 输出文档路径
            
        Returns:
            生成的文档路径
        """
        # 确保输出目录存在
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # 如果路径不是.docx结尾，添加.txt后缀
        if not output_path.lower().endswith('.docx'):
            output_path = f"{output_path}.txt"
        else:
            output_path = output_path.replace('.docx', '.txt')
        
        # 提取文本内容
        content = []
        
        # 添加标题
        content.append("# OCR识别结果文档")
        content.append(f"# 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")
        
        # 添加文本块
        for i, block in enumerate(ocr_result.get("structure", {}).get("text_blocks", [])):
            content.append(f"## 文本块 {i+1}")
            
            # 如果有原始文本和修正后的文本，显示对比
            if "original_text" in block and block["text"] != block["original_text"]:
                content.append("原始文本:")
                content.append(block["original_text"])
                content.append("")
                content.append("修正后文本:")
                content.append(block["text"])
            else:
                content.append(block["text"])
            
            content.append("")
        
        # 添加表格
        for i, table in enumerate(ocr_result.get("structure", {}).get("tables", [])):
            content.append(f"## 表格 {i+1}")
            content.append(table.get("content", ""))
            content.append("")
        
        # 添加图片引用
        for i, figure in enumerate(ocr_result.get("structure", {}).get("figures", [])):
            content.append(f"## 图片 {i+1}")
            content.append(f"[图片 {i+1}]")
            content.append("")
        
        # 添加元数据
        content.append("## 元数据")
        if "metadata" in ocr_result:
            for key, value in ocr_result["metadata"].items():
                content.append(f"{key}: {value}")
        
        # 添加Gemini修正信息
        if "gemini_corrections" in ocr_result:
            content.append("")
            content.append("## Gemini修正信息")
            corrections = ocr_result["gemini_corrections"]
            content.append(f"修正块数: {corrections.get('blocks_corrected', 0)}")
            content.append(f"处理时间: {corrections.get('processing_time', 0):.2f}秒")
            
            if "corrections" in corrections:
                content.append("")
                content.append("### 详细修正")
                for i, correction in enumerate(corrections["corrections"]):
                    content.append(f"#### 修正 {i+1}")
                    content.append(f"块ID: {correction.get('block_id', '')}")
                    content.append(f"原始文本: {correction.get('original_text', '')}")
                    content.append(f"修正文本: {correction.get('corrected_text', '')}")
                    content.append(f"修正前置信度: {correction.get('confidence_before', 0):.2f}")
                    content.append(f"修正后置信度: {correction.get('confidence_after', 0):.2f}")
                    content.append("")
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        logger.info(f"模拟Word文档生成完成: {output_path}")
        return output_path
    
    def _generate_real_document(self, ocr_result: Dict[str, Any], output_path: str) -> str:
        """
        使用python-docx生成实际的Word文档
        
        Args:
            ocr_result: OCR处理结果
            output_path: 输出文档路径
            
        Returns:
            生成的文档路径
        """
        # 确保输出目录存在
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # 如果路径不是.docx结尾，添加后缀
        if not output_path.lower().endswith('.docx'):
            output_path = f"{output_path}.docx"
        
        # 创建新文档
        doc = self.docx.Document()
        
        # 添加标题
        doc.add_heading('OCR识别结果文档', 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # 添加文本块
        for i, block in enumerate(ocr_result.get("structure", {}).get("text_blocks", [])):
            doc.add_heading(f"文本块 {i+1}", level=1)
            
            # 如果有原始文本和修正后的文本，显示对比
            if "original_text" in block and block["text"] != block["original_text"]:
                doc.add_heading("原始文本:", level=2)
                doc.add_paragraph(block["original_text"])
                
                doc.add_heading("修正后文本:", level=2)
                doc.add_paragraph(block["text"])
            else:
                doc.add_paragraph(block["text"])
            
            doc.add_paragraph()
        
        # 添加表格
        for i, table_info in enumerate(ocr_result.get("structure", {}).get("tables", [])):
            doc.add_heading(f"表格 {i+1}", level=1)
            
            # 这里应该解析HTML表格并创建Word表格
            # 由于这是简化版本，我们只添加表格内容的文本表示
            doc.add_paragraph(table_info.get("content", ""))
            doc.add_paragraph()
        
        # 添加图片引用
        for i, figure in enumerate(ocr_result.get("structure", {}).get("figures", [])):
            doc.add_heading(f"图片 {i+1}", level=1)
            doc.add_paragraph(f"[图片 {i+1}]")
            doc.add_paragraph()
        
        # 添加元数据
        doc.add_heading("元数据", level=1)
        if "metadata" in ocr_result:
            for key, value in ocr_result["metadata"].items():
                doc.add_paragraph(f"{key}: {value}")
        
        # 添加Gemini修正信息
        if "gemini_corrections" in ocr_result:
            doc.add_page_break()
            doc.add_heading("Gemini修正信息", level=1)
            corrections = ocr_result["gemini_corrections"]
            doc.add_paragraph(f"修正块数: {corrections.get('blocks_corrected', 0)}")
            doc.add_paragraph(f"处理时间: {corrections.get('processing_time', 0):.2f}秒")
            
            if "corrections" in corrections:
                doc.add_heading("详细修正", level=2)
                for i, correction in enumerate(corrections["corrections"]):
                    doc.add_heading(f"修正 {i+1}", level=3)
                    doc.add_paragraph(f"块ID: {correction.get('block_id', '')}")
                    doc.add_paragraph(f"原始文本: {correction.get('original_text', '')}")
                    doc.add_paragraph(f"修正文本: {correction.get('corrected_text', '')}")
                    doc.add_paragraph(f"修正前置信度: {correction.get('confidence_before', 0):.2f}")
                    doc.add_paragraph(f"修正后置信度: {correction.get('confidence_after', 0):.2f}")
                    doc.add_paragraph()
        
        # 保存文档
        doc.save(output_path)
        
        logger.info(f"Word文档生成完成: {output_path}")
        return output_path

# 测试代码
if __name__ == "__main__":
    # 创建生成器实例
    generator = WordGenerator(use_mock=True)
    
    # 创建测试OCR结果
    test_ocr_result = {
        "structure": {
            "text_blocks": [
                {
                    "id": "text_1",
                    "text": "这是一个测试文本，没有错误。",
                    "words": [
                        {"id": "word_1", "text": "这是", "confidence": 0.95},
                        {"id": "word_2", "text": "一个", "confidence": 0.92},
                        {"id": "word_3", "text": "测试", "confidence": 0.98},
                        {"id": "word_4", "text": "文本", "confidence": 0.97},
                        {"id": "word_5", "text": "，", "confidence": 0.99},
                        {"id": "word_6", "text": "没有", "confidence": 0.96},
                        {"id": "word_7", "text": "错误", "confidence": 0.94},
                        {"id": "word_8", "text": "。", "confidence": 0.99}
                    ]
                },
                {
                    "id": "text_2",
                    "original_text": "这个彳亍有一些未月显的错误。",
                    "text": "这个行走有一些明月显的错误。",
                    "words": [
                        {"id": "word_9", "text": "这个", "confidence": 0.93},
                        {"id": "word_10", "text": "行走", "confidence": 0.85, "original_text": "彳亍"},
                        {"id": "word_11", "text": "有", "confidence": 0.97},
                        {"id": "word_12", "text": "一些", "confidence": 0.95},
                        {"id": "word_13", "text": "明月", "confidence": 0.90, "original_text": "未月"},
                        {"id": "word_14", "text": "显", "confidence": 0.85},
                        {"id": "word_15", "text": "的", "confidence": 0.98},
                        {"id": "word_16", "text": "错误", "confidence": 0.96},
                        {"id": "word_17", "text": "。", "confidence": 0.99}
                    ]
                }
            ],
            "tables": [
                {
                    "id": "table_1",
                    "content": "<table><tr><td>项目</td><td>数量</td></tr><tr><td>测试项</td><td>5</td></tr></table>"
                }
            ],
            "figures": [
                {
                    "id": "figure_1"
                }
            ]
        },
        "metadata": {
            "file_path": "test_document.pdf",
            "processing_time": 1.25
        },
        "gemini_corrections": {
            "blocks_corrected": 1,
            "processing_time": 0.75,
            "corrections": [
                {
                    "block_id": "text_2",
                    "original_text": "这个彳亍有一些未月显的错误。",
                    "corrected_text": "这个行走有一些明月显的错误。",
                    "confidence_before": 0.75,
                    "confidence_after": 0.90
                }
            ]
        }
    }
    
    # 生成文档
    output_path = "test_output.docx"
    result_path = generator.generate_document(test_ocr_result, output_path)
    
    print(f"文档已生成: {result_path}")
